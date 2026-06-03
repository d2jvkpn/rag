import base64
import importlib.util
import io
import pathlib
import unittest


SCRIPT_PATH = pathlib.Path(__file__).parent.parent / "scripts" / "parse_docx.py"
SPEC = importlib.util.spec_from_file_location("parse_docx", SCRIPT_PATH)
parse_docx = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(parse_docx)

_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwADhQGAWjR9awAAAABJRU5ErkJggg=="
)


def _make_doc(*items):
    """Build a python-docx Document from a sequence of (kind, ...) tuples."""
    from docx import Document
    doc = Document()
    for item in items:
        kind = item[0]
        if kind == "heading":
            doc.add_heading(item[1], level=item[2] if len(item) > 2 else 1)
        elif kind == "paragraph":
            doc.add_paragraph(item[1])
        elif kind == "table":
            rows_data = item[1]
            t = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
            for r, row in enumerate(rows_data):
                for c, val in enumerate(row):
                    t.cell(r, c).text = val
        elif kind == "image":
            from docx.shared import Inches
            doc.add_picture(io.BytesIO(_PNG_1X1), width=Inches(1))
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


class ParseDocxSectionTests(unittest.TestCase):
    def test_single_section_no_heading(self):
        doc = _make_doc(("paragraph", "Hello world"))
        sections = parse_docx.build_sections(doc)
        self.assertEqual(len(sections), 1)
        texts = [i["text"] for i in sections[0]["items"] if i["kind"] == "text"]
        self.assertIn("Hello world", texts)

    def test_heading_splits_sections(self):
        doc = _make_doc(
            ("heading", "Intro", 1),
            ("paragraph", "First"),
            ("heading", "Details", 1),
            ("paragraph", "Second"),
        )
        sections = parse_docx.build_sections(doc)
        headings = [s["heading"] for s in sections if s["heading"]]
        self.assertIn("Intro", headings)
        self.assertIn("Details", headings)

    def test_cjk_heading_style_detected(self):
        """A paragraph with style val '标题1' or '1' should start a new section."""
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree

        doc = Document()
        p = doc.add_paragraph("CJK heading")
        # Inject pStyle with val='标题1'
        pPr = p._p.get_or_add_pPr()
        pStyle = etree.SubElement(pPr, qn("w:pStyle"))
        pStyle.set(qn("w:val"), "标题1")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)

        # _is_heading should recognise it
        body = doc2.element.body
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                pstyle = child.find(f".//{parse_docx._wtag('pStyle')}")
                if pstyle is not None:
                    val = pstyle.get(f"{{{parse_docx._W_NS}}}val") or ""
                    if val == "标题1":
                        self.assertTrue(parse_docx._is_heading(child))
                        return
        self.fail("paragraph with 标题1 style not found")

    def test_page_count_is_one(self):
        doc = _make_doc(
            ("heading", "A", 1),
            ("paragraph", "a"),
            ("heading", "B", 1),
            ("paragraph", "b"),
        )
        sections = parse_docx.build_sections(doc)
        page_data = []
        for section in sections:
            text, refs = parse_docx.render_section(section)
            if text or refs:
                page_data.append({"text": text, "page": section["section_num"], "refs": refs})
        # page_count is hardcoded to 1 in main(); verify sections exist but page_count is 1
        self.assertGreater(len(page_data), 1)
        # page_count used in payload
        import json, io as _io, sys, tempfile
        from docx import Document
        from pathlib import Path
        d = Document()
        d.add_heading("X", level=1)
        d.add_paragraph("p")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            d.save(f.name)
            tmp = f.name
        old_argv = sys.argv
        old_stdout = sys.stdout
        try:
            sys.argv = ["parse_docx.py", tmp]
            buf = _io.StringIO()
            sys.stdout = buf
            parse_docx.main()
            sys.stdout = old_stdout
            payload = json.loads(buf.getvalue())
            self.assertEqual(payload["page_count"], 1)
        finally:
            sys.argv = old_argv
            sys.stdout = old_stdout
            Path(tmp).unlink(missing_ok=True)


class ParseDocxTableTests(unittest.TestCase):
    def test_table_rendered_without_label_prefix(self):
        doc = _make_doc(
            ("heading", "Data", 1),
            ("table", [["Name", "Score"], ["Alice", "90"]]),
        )
        sections = parse_docx.build_sections(doc)
        section_with_data = next(s for s in sections if s["heading"] == "Data")
        text, refs = parse_docx.render_section(section_with_data)
        # Table rendered as raw Markdown, NO "Table N (Section N)" label prefix in text
        self.assertIn("| Name | Score |", text)
        self.assertNotIn("Table 1 (Section", text)
        # But the table ref carries the label
        table_ref = next((r for r in refs if r["ref_type"] == "table"), None)
        self.assertIsNotNone(table_ref)
        self.assertIn("Table 1", table_ref["label"])

    def test_adjacent_tables_merged(self):
        """Two consecutive tables with the same column count must be merged."""
        doc = _make_doc(
            ("heading", "Stats", 1),
            ("table", [["Metric", "Value"], ["Revenue", "100"]]),
            ("table", [["Metric", "Value"], ["Cost", "60"]]),
        )
        sections = parse_docx.build_sections(doc)
        section = next(s for s in sections if s["heading"] == "Stats")
        items = parse_docx._merge_adjacent_tables(section["items"])
        table_items = [i for i in items if i["kind"] == "table"]
        self.assertEqual(len(table_items), 1, "two adjacent tables should merge into one")
        rows = table_items[0]["rows"]
        # Header deduplicated
        header_rows = [r for r in rows if r[0] == "Metric"]
        self.assertEqual(len(header_rows), 1)
        texts = [r[0] for r in rows]
        self.assertIn("Revenue", texts)
        self.assertIn("Cost", texts)

    def test_adjacent_tables_different_cols_not_merged(self):
        doc = _make_doc(
            ("heading", "X", 1),
            ("table", [["A", "B"], ["1", "2"]]),
            ("table", [["C", "D", "E"], ["3", "4", "5"]]),
        )
        sections = parse_docx.build_sections(doc)
        section = next(s for s in sections if s["heading"] == "X")
        items = parse_docx._merge_adjacent_tables(section["items"])
        table_items = [i for i in items if i["kind"] == "table"]
        self.assertEqual(len(table_items), 2)


class ParseDocxImageTests(unittest.TestCase):
    def test_image_placeholder_in_body_text(self):
        doc = _make_doc(("image", None))
        sections = parse_docx.build_sections(doc)
        all_items = [i for s in sections for i in s["items"]]
        text_items = [i for i in all_items if i["kind"] == "text"]
        placeholder_texts = [i["text"] for i in text_items if "[Image:" in i["text"]]
        self.assertTrue(len(placeholder_texts) > 0, "expected [Image:...] placeholder in text")
        ref_ids = [r["ref_id"] for i in text_items for r in i.get("refs", [])
                   if r["ref_type"] == "image"]
        self.assertTrue(len(ref_ids) > 0)

    def test_image_ref_without_storage(self):
        doc = _make_doc(("image", None))
        sections = parse_docx.build_sections(doc)
        section_with_image = next(s for s in sections if any(
            r["ref_type"] == "image"
            for i in s["items"] for r in i.get("refs", [])
        ))
        sn = section_with_image["section_num"]
        text, refs = parse_docx.render_section(section_with_image)
        image_ref = next((r for r in refs if r["ref_type"] == "image"), None)
        self.assertIsNotNone(image_ref)
        self.assertTrue(image_ref["ref_id"].startswith(f"docx-image-{sn}-"))
        self.assertNotIn("storage_path", image_ref)

    def test_image_ref_with_storage(self):
        import tempfile
        doc = _make_doc(("image", None))

        with tempfile.TemporaryDirectory() as tmp:
            storage_base = pathlib.Path(tmp)
            media_dir = storage_base / "doc-1"
            sections = parse_docx.build_sections(doc, media_dir, storage_base)
            section_with_image = next(s for s in sections if any(
                r["ref_type"] == "image"
                for i in s["items"] for r in i.get("refs", [])
            ))
            text, refs = parse_docx.render_section(
                section_with_image, media_dir, storage_base
            )
            image_ref = next((r for r in refs if r["ref_type"] == "image"), None)
            self.assertIsNotNone(image_ref)
            self.assertIn("storage_path", image_ref)
            self.assertTrue((storage_base / image_ref["storage_path"]).is_file())


if __name__ == "__main__":
    unittest.main()
